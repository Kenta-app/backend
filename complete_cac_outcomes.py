from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:/Users/sdiaz/Downloads/ACP_Diaz_Salvador_Cama_Jimena.docx")
OUTPUT = Path(r"C:/Users/sdiaz/Downloads/ACP_Diaz_Salvador_Cama_Jimena_SO_CAC_completado.docx")


outcomes = [
    (
        "STUDENT OUTCOME CAC 1",
        "La capacidad de analizar un problema complejo y aplicar principios de computación y otras disciplinas relevantes para identificar soluciones.",
        "En el proyecto Kenta, este resultado se evidencia en el análisis de la saturación informativa, la desinformación política y la baja confianza institucional que afectan a los jóvenes de Lima Metropolitana. Para abordar este problema complejo, se aplican principios de computación como procesamiento de lenguaje natural, aprendizaje automático, clasificación de textos, detección de noticias falsas, análisis de sentimiento y diseño de sistemas web. Además, se integran conocimientos de ciencia política, comunicación digital y educación cívica para comprender el contexto social del usuario y definir una solución pertinente.",
        "La solución propuesta parte de un benchmarking de modelos, frameworks y fuentes de información; luego identifica alternativas técnicas viables para recopilar, procesar, resumir y presentar contenido político confiable. De esta manera, el equipo transforma una necesidad social amplia en requerimientos técnicos concretos, medibles y alineados con los objetivos del proyecto.",
        "Evidencias: análisis del problema, benchmarking tecnológico, documento de requerimientos, modelo de datos, criterios de selección de modelos ML/NLP y matriz de riesgos asociados a datos, sesgos e integración.",
    ),
    (
        "STUDENT OUTCOME CAC 2",
        "La capacidad de diseñar, implementar y evaluar una solución basada en computación para cumplir con el conjunto de requerimientos en el contexto de sistemas de información.",
        "Este resultado se cumple mediante el diseño e implementación de una aplicación web basada en Machine Learning y Procesamiento de Lenguaje Natural que responde a requerimientos funcionales y técnicos definidos para el sistema de información. Kenta contempla un frontend responsivo, un backend en Python, una base de datos para usuarios, preferencias e historial, APIs documentadas y módulos especializados para autenticación, procesamiento de texto, clasificación temática, generación de resúmenes y recomendaciones.",
        "La evaluación de la solución considera métricas de desempeño del sistema y de los modelos, como precisión mínima del 80% para clasificación, métricas ROUGE y BERTScore para resumen automático, cobertura de pruebas unitarias mayor o igual al 80%, disponibilidad mayor al 99% y tiempos de respuesta definidos para escenarios de carga estándar. Con ello, el proyecto no solo propone una arquitectura, sino que también establece criterios verificables para comprobar si la solución cumple los requerimientos del contexto.",
        "Evidencias: prototipo UX/UI, especificación de APIs, arquitectura del sistema, implementación del backend y frontend, modelo de ML entrenado, módulo NLP, pruebas unitarias, pruebas funcionales y reporte de rendimiento.",
    ),
    (
        "STUDENT OUTCOME CAC 3",
        "La capacidad de comunicarse efectivamente con un rango de audiencias y variedad de contextos profesionales.",
        "El proyecto requiere comunicación efectiva con diferentes audiencias: comité de proyectos, portfolio manager, product owner, equipo técnico y usuarios jóvenes de Lima Metropolitana. Para ello, el equipo documenta el Project Charter, los objetivos, indicadores, alcance, riesgos, arquitectura, backlog, avances por sprint y resultados de validación en un lenguaje adecuado para audiencias técnicas y no técnicas.",
        "Asimismo, Kenta exige traducir información política compleja en resúmenes claros, comprensibles y confiables para usuarios jóvenes. Esta capacidad comunicativa se refleja tanto en la gestión del proyecto como en el propio diseño del producto, ya que la interfaz debe presentar resultados de clasificación, recomendaciones y resúmenes de manera accesible, evitando lenguaje excesivamente técnico y facilitando la comprensión del contenido político.",
        "Evidencias: Project Charter, actas de reunión, sustentaciones, documentación técnica, prototipo UX/UI, historias de usuario, presentación de avances y reportes de validación con stakeholders.",
    ),
    (
        "STUDENT OUTCOME CAC 4",
        "La capacidad de reconocer responsabilidades éticas y profesionales en situaciones de ingeniería y hacer juicios informados, que deben considerar el impacto de las soluciones de ingeniería en contextos globales, económicos, ambientales y sociales.",
        "Kenta trabaja con información política, por lo que el equipo debe reconocer responsabilidades éticas relacionadas con la transparencia, el tratamiento responsable de datos, el sesgo algorítmico, la manipulación informativa y el impacto social de los sistemas de recomendación. El proyecto considera que una solución de IA aplicada a política no debe reemplazar el juicio crítico del usuario, sino facilitar el acceso a información resumida, clasificada y comprensible.",
        "Desde el punto de vista profesional, se establecen límites claros en el alcance: el sistema no realiza auditoría definitiva del contenido político externo ni predicción electoral. Además, se contempla la mitigación de riesgos como datasets sesgados, resúmenes incoherentes, bajo desempeño en detección de noticias falsas y dependencia de herramientas externas. Estos juicios consideran el impacto social de la plataforma en la participación cívica juvenil, así como restricciones económicas y tecnológicas del entorno académico.",
        "Evidencias: alcance y exclusiones del proyecto, análisis de riesgos, criterios de calidad de datos, validación de modelos, documentación de sesgos, selección responsable de fuentes públicas y definición de límites funcionales del sistema.",
    ),
    (
        "STUDENT OUTCOME CAC 5",
        "La capacidad de funcionar efectivamente como miembro o lider en un equipo cuyos miembros juntos proporcionan liderazgo, crean un entorno de colaboración e inclusivo, establecen objetivos, planifican tareas y cumplen objetivos.",
        "Este resultado se evidencia en la organización del proyecto bajo roles definidos de Project Manager y Scrum Master, así como en la aplicación de Scrum para planificar, ejecutar y revisar el trabajo mediante sprints. El equipo establece objetivos generales y específicos, prioriza funcionalidades en el backlog, define hitos académicos, distribuye responsabilidades y realiza seguimiento de entregables para cumplir con los plazos establecidos.",
        "La colaboración se manifiesta en la coordinación continua entre los miembros del equipo, la revisión de avances, la resolución de impedimentos y la comunicación con stakeholders. El liderazgo se comparte entre la gestión del alcance, la facilitación de la metodología ágil y la toma de decisiones técnicas necesarias para desarrollar una solución de software con componentes de frontend, backend, base de datos y modelos de inteligencia artificial.",
        "Evidencias: organigrama del proyecto, roles y responsabilidades, plan de trabajo, product backlog, actas de reunión, cronograma de hitos, seguimiento de sprints y entregables aprobados.",
    ),
    (
        "STUDENT OUTCOME CAC 6",
        "La capacidad para comprender y brindar soporte para el uso, entrega y gestión de sistemas de información dentro de un entorno de sistemas de información.",
        "Kenta se desarrolla como un sistema de información que integra usuarios, datos, procesos, modelos de análisis y servicios web para entregar información política procesada. Este resultado se evidencia en la comprensión del ciclo de uso del sistema: ingreso o recopilación de contenido político, preprocesamiento de texto, clasificación, resumen automático, almacenamiento de resultados, recomendación personalizada y presentación al usuario final mediante una interfaz web.",
        "El soporte para la entrega y gestión del sistema se refleja en la documentación técnica, la definición de arquitectura, el diseño de APIs, la gestión de base de datos, las pruebas funcionales y de rendimiento, y la planificación del despliegue bajo restricciones de infraestructura. Además, el proyecto contempla la experiencia de usuarios jóvenes, la disponibilidad del servicio, la conectividad y la mantenibilidad básica de los componentes desarrollados.",
        "Evidencias: arquitectura del sistema, documentación de APIs, modelo de base de datos, manual o guía técnica de despliegue, pruebas funcionales, pruebas de rendimiento, gestión de usuarios y documentación del flujo de procesamiento ML/NLP.",
    ),
]


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    if not value:
        keep.set(qn("w:val"), "0")


def set_run_font(run, bold=False, size=11):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def style_body(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        set_run_font(run, size=11)


doc = Document(SOURCE)

anchor = None
for paragraph in doc.paragraphs:
    if paragraph.text.strip().lower() == "riesgos y mitigación":
        anchor = paragraph
        break

if anchor is None:
    raise RuntimeError("No se encontró el encabezado 'Riesgos y Mitigación'.")

inserted = []
inserted.append(insert_paragraph_before(anchor, "", None))
title = insert_paragraph_before(anchor, "Student Outcomes CAC", "Heading 1")
set_keep_with_next(title)
inserted.append(title)

intro = insert_paragraph_before(
    anchor,
    "A continuación, se describe cómo el proyecto Kenta evidencia el cumplimiento de los Student Outcomes CAC, considerando el análisis, diseño, implementación, gestión y validación de una solución basada en sistemas de información, Machine Learning y Procesamiento de Lenguaje Natural.",
    "Normal",
)
style_body(intro)
inserted.append(intro)

for code, definition, application, second, evidence in outcomes:
    heading = insert_paragraph_before(anchor, code, "Heading 2")
    set_keep_with_next(heading)
    inserted.append(heading)

    def_p = insert_paragraph_before(anchor, "", "Normal")
    r = def_p.add_run("Resultado esperado: ")
    set_run_font(r, bold=True, size=11)
    r = def_p.add_run(definition)
    set_run_font(r, size=11)
    style_body(def_p)
    inserted.append(def_p)

    for text in (application, second):
        p = insert_paragraph_before(anchor, text, "Normal")
        style_body(p)
        inserted.append(p)

    ev_p = insert_paragraph_before(anchor, "", "Normal")
    r = ev_p.add_run("Evidencias: ")
    set_run_font(r, bold=True, size=11)
    ev_text = evidence.removeprefix("Evidencias: ")
    r = ev_p.add_run(ev_text)
    set_run_font(r, size=11)
    style_body(ev_p)
    inserted.append(ev_p)

    spacer = insert_paragraph_before(anchor, "", "Normal")
    spacer.paragraph_format.space_after = Pt(2)
    inserted.append(spacer)

doc.save(OUTPUT)
print(OUTPUT)
